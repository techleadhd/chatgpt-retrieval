import os
import sys
# new code
from docx import Document
from funcoms import get_loader_cls
# end new code

import openai
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import TextLoader, PDFMinerLoader
from langchain.embeddings import OpenAIEmbeddings
from langchain.indexes import VectorstoreIndexCreator
from langchain.indexes.vectorstore import VectorStoreIndexWrapper
from langchain.llms import OpenAI
from langchain.vectorstores import Chroma
# new code
#from langchain.docstore import Document

class WordLoader:
    def __init__(self, filepath):
        self.filepath = filepath

    def load(self):
        doc = docx.Document(self.filepath)
        for paragraph in doc.paragraphs:
            yield Document(text=paragraph.text)


""" def get_loader_cls(file_type: str):
    if file_type == '.txt':
        return TextLoader
    elif file_type == '.pdf':
        return PDFMinerLoader
    elif file_type == '.docx':
        return WordLoader
    else:
        raise ValueError(f'No loader found for file type {file_type}')
 """


#new code ends
import constants
data_directory = constants.DATA_DIRECTORY

os.environ["OPENAI_API_KEY"] = constants.APIKEY

# Enable to cache & reuse the model to disk (for repeated queries on the same data)
PERSIST = constants.PERSIST #False

query = sys.argv[1]

if PERSIST and os.path.exists("persist"):
  print("Reusing index...\n")
  vectorstore = Chroma(persist_directory="persist", embedding_function=OpenAIEmbeddings())
  index = VectorStoreIndexWrapper(vectorstore=vectorstore)
else:
  # new code starts
    index_creator = VectorstoreIndexCreator(vectorstore_kwargs={"persist_directory":"persist"} if PERSIST else {})
    loaders = []
    for filename in os.listdir( data_directory ):
        file_extension = os.path.splitext(filename)[-1]
        loader_cls = get_loader_cls(file_extension)
        if loader_cls is not None:  # Ensure a loader exists for this file type
            loader = loader_cls(os.path.join( data_directory , filename))
            loaders.append(loader)
    if loaders:
        index = index_creator.from_loaders(loaders)
    else:
        print("No valid files found in the data directory.")  
  #loader = TextLoader('mydata.txt')
  # This code can also import folders, including various filetypes like PDFs using the DirectoryLoader.
  #loader = DirectoryLoader("./mydata/", glob="*.txt",loader_cls=TextLoader)
  #if PERSIST:
  #    index = VectorstoreIndexCreator(vectorstore_kwargs={"persist_directory":"persist"}).from_loaders([loader])    
  #else:
  #  index = VectorstoreIndexCreator().from_loaders([loader])
  # new code ends

chain = RetrievalQA.from_chain_type(
  llm=ChatOpenAI(model="gpt-3.5-turbo"),
  retriever=index.vectorstore.as_retriever(search_kwargs={"k": 1}),
)
print(chain.run(query))
