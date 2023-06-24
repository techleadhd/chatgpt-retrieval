
from docx import Document

class WordLoader:
    def __init__(self, filepath):
        self.filepath = filepath

    def load(self):
        doc = docx.Document(self.filepath)
        for paragraph in doc.paragraphs:
            yield Document(text=paragraph.text)


def get_loader_cls(file_extension: str):
    from langchain.document_loaders import TextLoader, PDFMinerLoader, BSHTMLLoader  # import any other loaders you might need

    if file_extension == '.txt':
        return TextLoader
    elif file_extension == '.pdf':
        return PDFMinerLoader
    elif file_extension == '.html':
        return BSHTMLLoader
    elif file_extension == '.docx':
        return WordLoader
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")
